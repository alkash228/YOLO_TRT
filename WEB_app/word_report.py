"""Word (.docx) incident report per violator stable_id."""
from __future__ import annotations

from dataclasses import replace
from datetime import datetime
from pathlib import Path
from typing import Any

import cv2
import numpy as np

from app.core.frame_pipeline import FramePacket
from app.core.video_encode import _EncodeJob, _RenderContext, _render_encode_job, resolve_run_packets
from WEB_app.video_builder import (
    _verdict_is_violation,
    filter_person_single_id,
    filter_violator_single_id,
    highlight_no_helmet_on_rgb,
    iter_run_packets,
    load_run_metadata,
    resolve_overlay,
    stamp_all_people_boxes_rgb,
    stamp_all_poses_rgb,
    stamp_debug_hud_rgb,
    stamp_helmet_accessories_rgb,
    web_debug_show_all_dets,
)

TEST_COMPANIES: tuple[str, ...] = (
    'ООО «Обнал»',
    'ООО «Рога и копыта»',
)

_MONTHS_RU = (
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)


def format_incident_datetime_ru(value: datetime) -> str:
    return (
        f"{value.day} {_MONTHS_RU[value.month - 1]} {value.year} г., "
        f"{value.hour:02d}:{value.minute:02d}"
    )


def parse_incident_datetime(raw: str) -> datetime:
    text = str(raw or "").strip()
    if not text:
        return datetime.now()
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        return datetime.fromisoformat(text)
    except ValueError:
        pass
    for fmt in ("%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M", "%d.%m.%Y %H:%M"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return datetime.now()


def _scan_violator_packets(
    data: dict[str, Any],
    run_dir: Path,
    stable_id: int,
    *,
    stop_after_first_violation: bool = False,
) -> tuple[list[int], list[int], FramePacket | None, int]:
    """
    One pass over spill packets:
    presence frames, violation frames, a packet for the chosen report frame, counts.
    Photo is the first NO HELMET frame (not mid-clip).
    """
    sid = int(stable_id)
    presence: list[int] = []
    violations: list[int] = []
    packets_by_frame: dict[int, FramePacket] = {}
    for packet in iter_run_packets(data, Path(run_dir)):
        n = int(packet.n_inst)
        if n <= 0 or packet.stable_ids is None:
            continue
        hit = False
        viol = False
        for i in range(min(n, len(packet.stable_ids))):
            if int(packet.stable_ids[i]) != sid:
                continue
            hit = True
            verdicts = packet.cross_check_verdicts or []
            if i < len(verdicts) and _verdict_is_violation(verdicts[i]):
                viol = True
            break
        if not hit:
            continue
        fi = int(packet.frame_idx)
        presence.append(fi)
        packets_by_frame[fi] = packet
        if viol:
            violations.append(fi)
            if stop_after_first_violation:
                break

    if violations:
        # Earliest NO HELMET — much faster on long HEVC than mid-video frame.
        frame_idx = int(violations[0])
    elif presence:
        frame_idx = int(presence[0])
    else:
        return [], [], None, -1

    return presence, violations, packets_by_frame.get(frame_idx), frame_idx


def _clip_path_for_id(run_dir: Path, run_id: str, stable_id: int) -> Path:
    """Exact clip path only — never glob (id=1 must not pick id=10 / other runs)."""
    return Path(run_dir) / f"{run_id}_violation_id{int(stable_id)}.mp4"


def _read_still_from_violation_clip(run_dir: Path, run_id: str, stable_id: int) -> np.ndarray | None:
    """
    Fast path: ONLY `{run_id}_violation_id{sid}.mp4` (exact name).
    Loose globs previously could feed the same wrong still into every report.
    """
    clip = _clip_path_for_id(run_dir, run_id, stable_id)
    if not clip.is_file() or clip.stat().st_size <= 0:
        return None
    cap = cv2.VideoCapture(str(clip))
    if not cap.isOpened():
        return None
    try:
        ok, frame = cap.read()
        if ok and frame is not None and frame.size > 0:
            return frame
    finally:
        cap.release()
    return None


def _annotate_stable_id(rgb: np.ndarray, stable_id: int, frame_idx: int) -> np.ndarray:
    """Burn ID into the still so reports for different IDs cannot be confused."""
    out = np.ascontiguousarray(rgb.copy())
    bgr = cv2.cvtColor(out, cv2.COLOR_RGB2BGR)
    label = f"ID {int(stable_id)}  |  frame {int(frame_idx)}"
    cv2.rectangle(bgr, (8, 8), (8 + 12 * len(label) + 24, 48), (0, 0, 0), -1)
    cv2.putText(
        bgr,
        label,
        (16, 38),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (0, 255, 255),
        2,
        cv2.LINE_AA,
    )
    return cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)

def read_source_frame_bgr(input_path: str, frame_idx: int) -> np.ndarray | None:
    """Grab one source frame. Sequential decode — HEVC-safe (no POS_FRAMES seek)."""
    from app.core.window_frame_loader import read_frame_bgr_sequential

    if not input_path:
        return None
    return read_frame_bgr_sequential(str(input_path), int(frame_idx))


def _imwrite_bgr(path: Path, bgr: np.ndarray, *, quality: int = 92) -> None:
    """cv2.imwrite breaks on non-ASCII Windows paths — encode then write bytes."""
    ok, buf = cv2.imencode(".jpg", bgr, [int(cv2.IMWRITE_JPEG_QUALITY), int(quality)])
    if not ok:
        raise RuntimeError("Не удалось закодировать JPEG для отчёта")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(buf.tobytes())


def _blank_bgr(width: int, height: int) -> np.ndarray:
    h = max(1, int(height) or 720)
    w = max(1, int(width) or 1280)
    return np.zeros((h, w, 3), dtype=np.uint8)


def render_person_frame_rgb(
    *,
    packet: FramePacket | None,
    stable_id: int,
    meta: dict[str, Any],
    overlay: dict[str, Any],
    frame_idx: int,
    prefer_violation: bool = True,
    run_dir: Path | None = None,
    run_id: str | None = None,
    source_override: str | None = None,
) -> np.ndarray:
    from app.core.video_encode import resolve_run_source_video

    recorded = str(meta.get("input_path") or meta.get("recorded_input_path") or "")
    input_path = recorded
    if run_dir is not None:
        resolved = resolve_run_source_video(
            Path(run_dir),
            recorded,
            run_id=run_id,
            override=source_override,
        )
        if resolved:
            input_path = resolved

    width = int(meta.get("width") or 0)
    height = int(meta.get("height") or 0)
    frame_bgr = read_source_frame_bgr(input_path, frame_idx) if input_path else None

    if packet is None:
        if frame_bgr is None:
            raise RuntimeError(
                f"Кадр {frame_idx} недоступен: нет packet и не читается исходник "
                f"({input_path or 'input_path пуст'}). Положи RUNID_source.mp4 в папку прогона."
            )
        return cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)

    if frame_bgr is None and packet.frame_bgr is not None and packet.frame_bgr.size > 0:
        fb = packet.frame_bgr
        if fb.shape[0] > 2 and fb.shape[1] > 2:
            frame_bgr = fb

    if frame_bgr is None:
        # Still draw overlay on a blank canvas if source video is missing on this machine.
        if packet.mask_hw:
            height, width = int(packet.mask_hw[0]), int(packet.mask_hw[1])
        frame_bgr = _blank_bgr(width, height)

    height, width = int(frame_bgr.shape[0]), int(frame_bgr.shape[1])

    person = filter_person_single_id(packet, stable_id)
    if person.n_inst <= 0:
        raise RuntimeError(
            f"На кадре {frame_idx} нет объекта ID {stable_id} в packets "
            "(нельзя подставлять чужой full-frame)."
        )

    debug_all = web_debug_show_all_dets()
    if debug_all:
        # Full scene for debug; still require this ID to be present.
        draw_packet = packet
    else:
        filtered = (
            filter_violator_single_id(packet, stable_id)
            if prefer_violation
            else person
        )
        if filtered.n_inst <= 0:
            filtered = person
        draw_packet = filtered

    prompt = str(meta.get("prompt") or "person").strip().casefold() or "person"
    ctx = _RenderContext(
        target_w=max(1, width),
        target_h=max(1, height),
        prompt_lookup={prompt: 0},
        overlay=overlay,
    )
    job = _EncodeJob(src_i=int(frame_idx), carry=draw_packet, frame_bgr=frame_bgr)
    _, rgb = _render_encode_job(job, ctx)
    pkt = replace(draw_packet, frame_idx=int(frame_idx))
    # Stills prioritize person + helmet boxes; pose only if enabled and keypoints exist.
    rgb = stamp_all_people_boxes_rgb(
        rgb,
        pkt,
        target_w=ctx.target_w,
        target_h=ctx.target_h,
        focus_sid=int(stable_id),
    )
    rgb = stamp_helmet_accessories_rgb(
        rgb,
        pkt,
        target_w=ctx.target_w,
        target_h=ctx.target_h,
    )
    if bool(overlay.get("draw_pose", False)):
        rgb = stamp_all_poses_rgb(
            rgb,
            pkt,
            target_w=ctx.target_w,
            target_h=ctx.target_h,
            pose_kpt_conf=float(overlay.get("pose_kpt_conf", 0.25)),
            pose_point_radius=int(overlay.get("pose_point_radius", 6) or 6),
            pose_line_thickness=int(overlay.get("pose_line_thickness", 3) or 3),
        )
    rgb = highlight_no_helmet_on_rgb(
        rgb,
        pkt,
        int(stable_id),
        target_w=ctx.target_w,
        target_h=ctx.target_h,
        frame_bgr=frame_bgr,
        force=True,
    )
    if debug_all:
        rgb = stamp_debug_hud_rgb(
            rgb,
            pkt,
            draw_pose=bool(overlay.get("draw_pose", False)),
        )
    return _annotate_stable_id(rgb, int(stable_id), int(frame_idx))


def build_word_report(
    run_dir: str | Path,
    run_id: str,
    stable_id: int,
    *,
    company: str,
    organization: str,
    incident_datetime: datetime,
    overlay_override: dict[str, Any] | None = None,
    source_video: str | None = None,
    violation_count: int | None = None,
    presence_frames: int | None = None,
) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError(
            "Установите python-docx: pip install python-docx"
        ) from exc

    run_path = Path(run_dir)
    sid = int(stable_id)
    meta = load_run_metadata(run_path, run_id, source_video=source_video)
    data, _ = resolve_run_packets(run_path, run_id=run_id)
    overlay = resolve_overlay(meta, overlay_override=overlay_override)

    presence, violations, packet, frame_idx = _scan_violator_packets(
        data, run_path, sid
    )
    if frame_idx < 0 or packet is None:
        raise RuntimeError(f"Нет кадров с нарушителем ID {sid}")

    # Always render from packets filtered by THIS stable_id.
    # Do not reuse clip stills — old/wrong clips produced the same photo for every ID.
    rgb = render_person_frame_rgb(
        packet=packet,
        stable_id=sid,
        meta=meta,
        overlay=overlay,
        frame_idx=frame_idx,
        prefer_violation=bool(violations),
        run_dir=run_path,
        run_id=run_id,
        source_override=source_video,
    )

    viol_count = int(violation_count) if violation_count is not None else len(violations)
    presence_count = int(presence_frames) if presence_frames is not None else len(presence)

    reports_dir = run_path / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    image_path = reports_dir / f"{run_id}_id{sid}_frame{frame_idx}.jpg"
    docx_path = reports_dir / f"{run_id}_akt_id{sid}.docx"

    bgr = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    try:
        _imwrite_bgr(image_path, bgr, quality=92)
    except OSError as exc:
        raise RuntimeError(f"Не удалось сохранить кадр для отчёта: {image_path} ({exc})") from exc

    org = (organization or company or "").strip() or "—"
    company_name = (company or org).strip() or "—"
    incident_text = format_incident_datetime_ru(incident_datetime)
    source_name = Path(str(meta.get("input_path") or "")).name or "—"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ фиксации нарушения\nтребований охраны труда")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    def add_row(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    add_row("Организация", org)
    add_row("Объект / заказчик (тест)", company_name)
    add_row("Дата и время инцидента", incident_text)
    add_row("Вид нарушения", "Работа без защитной каски (NO HELMET)")
    add_row("Идентификатор нарушителя", f"№ {sid}")
    add_row("Кадров присутствия в видео", str(presence_count))
    add_row("Срабатываний без каски", str(viol_count))
    add_row("Кадр в отчёте", str(frame_idx))
    add_row("Исходное видео", source_name)

    doc.add_paragraph()
    cap = doc.add_paragraph("Фотофиксация нарушителя:")
    cap.runs[0].bold = True
    from io import BytesIO

    doc.add_picture(BytesIO(image_path.read_bytes()), width=Cm(14))

    doc.save(str(docx_path))
    return docx_path
